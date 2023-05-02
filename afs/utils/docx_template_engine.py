from base64 import b64decode
import copy
from io import BytesIO
import re
from typing import List, Tuple

from django.utils.translation import ugettext as _
from afs.utils.arches_template_engine import ArchesTemplateEngine, TemplateTag, TemplateTagType
import docx
from docx.oxml.ns import qn
from docx.document import Document as _Document
from docx.oxml import OxmlElement
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.section import _Header


class DocxTemplateEngine(ArchesTemplateEngine):
    def extract_regex_matches(self, template) -> List[Tuple]:
        self.doc = docx.Document(template)
        tags = self.iterate_over_container(self.doc)

        return tags

    def iterate_over_container(self, container, parent=None):
        parsed_tags: List[Tuple] = []
        try:
            for section in container.sections:
                parsed_tags += self.iterate_over_container(section.header)
        except AttributeError:
            pass  # this is ok, there are lots of types that do not have a "sections" attribute - skip them and continue

        for block in self.iter_block_items(container):
            if isinstance(block, Paragraph):
                for match in re.findall(self.regex, block.text):
                    parsed_tags.append((match, {"docxBlock": block, "parent": parent}))

            elif isinstance(block, Table):
                row_length = len(block.rows)
                column_length = len(block.columns)
                current_row = 0
                while current_row < row_length:
                    current_column = 0
                    while current_column < column_length:
                        current_cell = block.cell(current_row, current_column)
                        parsed_tags += self.iterate_over_container(current_cell, block)
                        current_column += 1
                    current_row += 1
                pass
        return parsed_tags

    def iter_block_items(self, parent):
        """
        Generate a reference to each paragraph and table child within *parent*,
        in document order. Each returned value is an instance of either Table or
        Paragraph. *parent* would most commonly be a reference to a main
        Document object, but also works for a _Cell object, which itself can
        contain paragraphs and tables.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
            element_parent = parent._body
        elif isinstance(parent, _Cell):
            element_parent = parent_elm = parent._tc
        elif isinstance(parent, _Header):
            element_parent = parent_elm = parent._element
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, element_parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, element_parent)

    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    def remove_row(table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

    def replace_tags(self, tags: List[TemplateTag]):
        for tag in tags:
            block = tag.optional_keys["docxBlock"]
            if tag.type == TemplateTagType.CONTEXT:
                if tag.has_rows:
                    column = 0
                    # this is ugly, but way more efficient than the alternative
                    parent_table = tag.context_children_template[-1].optional_keys["parent"]
                    current_row = parent_table.add_row()

                    for child in tag.children:
                        if child.type == TemplateTagType.ROWEND:
                            column = -1
                            current_row = parent_table.add_row()
                        elif child.type == TemplateTagType.VALUE:
                            # grab any borders from the original cell copy them to the new cell.
                            template_block = tag.context_children_template[column].optional_keys["docxBlock"]
                            borders = template_block._parent.get_or_add_tcPr().first_child_found_in("w:tcBorders")

                            for edge in ("start", "top", "end", "bottom", "left", "right", "insideH", "insideV"):
                                raw_border_tag = "w:{}".format(edge)

                                # check for tag existnace, if none found, then create one
                                element = borders.find(qn(raw_border_tag))
                                cell_borders = current_row.cells[column]._tc.get_or_add_tcPr().first_child_found_in("w:tcBorders")
                                if cell_borders is None:
                                    cell_borders = OxmlElement("w:tcBorders")
                                    current_row.cells[column]._tc.get_or_add_tcPr().append(cell_borders)
                                if element is not None:
                                    cell_borders.append(copy.deepcopy(element))
                            # every cell gets created with (bad) default styling.
                            DocxTemplateEngine.delete_paragraph(current_row.cells[column].paragraphs[0])
                            # copies paragraph styling from the original template cells
                            current_row.cells[column].add_paragraph(
                                "" if child.value == None else child.value,
                                copy.deepcopy(_Cell(template_block._parent, parent_table).paragraphs[0].style),
                            )
                        column += 1

                    if tag.attributes["has_header"] == "true":
                        DocxTemplateEngine.remove_row(parent_table, parent_table.rows[1])
                    else:
                        DocxTemplateEngine.remove_row(parent_table, parent_table.rows[0])
                else:
                    self.replace_tags(tag.children)
                DocxTemplateEngine.delete_paragraph(block)
                DocxTemplateEngine.delete_paragraph(tag.end_tag.optional_keys["docxBlock"])
            elif tag.type == TemplateTagType.VALUE:
                block.text = tag.value
            elif tag.type == TemplateTagType.IMAGE:
                block.text = ""
                run = block.add_run()
                if tag.value:
                    run.add_picture(BytesIO(b64decode(re.sub("data:image/jpeg;base64,", "", tag.value))))

    def create_file(self, tags: List[TemplateTag], template):
        bytestream = BytesIO()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        self.replace_tags(tags)
        self.doc.save(bytestream)
        return (bytestream, mime, "test.docx")
